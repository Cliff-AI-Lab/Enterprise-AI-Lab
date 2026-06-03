"""会议纪要导出(PRD 4.1.4):从已落库事件重建结构化纪要 → Markdown / DOCX / PDF。

纪要至少含:议程回顾、各议题讨论摘要、最终结论、风险清单、行动建议(AC2/AC3)。
PDF 用 reportlab + macOS 系统中文字体(无系统库依赖)。
"""
from __future__ import annotations

import io
import os

from . import store

# macOS 常见中文字体候选(TTC 用 subfontIndex=0)
_CJK_FONTS = [
    "/System/Library/Fonts/Supplemental/Songti.ttc",
    "/System/Library/Fonts/STHeiti Light.ttc",
    "/System/Library/Fonts/PingFang.ttc",
    "/Library/Fonts/Arial Unicode.ttf",
]


def build_struct(mid: str) -> dict:
    """把事件流重建成纪要结构。"""
    m = store.get_meeting(mid)
    if not m:
        raise ValueError("会议不存在")
    experts = store.get_experts_by_ids(m["expert_ids"])
    events = store.get_events(mid)
    agenda = m["agenda"] or []
    topics: dict[int, dict] = {}
    conclusion = None
    opening = ""
    for ev in events:
        t = ev.get("type")
        if t == "moderator.turn":
            if ev.get("phase") == "opening":
                opening = ev.get("text", "")
            else:
                ti = ev.get("topic_index", 0)
                tp = topics.setdefault(ti, {"title": agenda[ti] if ti < len(agenda) else "", "turns": [], "conflict": None, "summary": None})
                tp["turns"].append({"name": "主持人", "domain": "主持", "text": ev.get("text", ""),
                                    "followup": False, "human": False})
        elif t == "topic.started":
            topics.setdefault(ev["index"], {"title": ev.get("title", ""), "turns": [], "conflict": None, "summary": None})
        elif t in ("turn", "human.turn"):
            ti = ev.get("topic_index", 0)
            tp = topics.setdefault(ti, {"title": agenda[ti] if ti < len(agenda) else "", "turns": [], "conflict": None, "summary": None})
            tp["turns"].append({
                "name": ev.get("expert_name") or ev.get("name", "用户"),
                "domain": ev.get("domain", "参会者"),
                "text": ev.get("text", ""),
                "followup": ev.get("phase") == "followup",
                "human": t == "human.turn",
            })
        elif t == "conflict":
            ti = ev.get("topic_index", 0)
            topics.setdefault(ti, {"title": "", "turns": [], "conflict": None, "summary": None})["conflict"] = ev
        elif t == "stage_summary":
            ti = ev.get("topic_index", 0)
            topics.setdefault(ti, {"title": "", "turns": [], "conflict": None, "summary": None})["summary"] = ev
        elif t == "conclusion":
            conclusion = ev
    return {"meeting": m, "experts": experts, "agenda": agenda, "opening": opening,
            "topics": [topics[k] for k in sorted(topics)], "conclusion": conclusion}


def to_markdown(s: dict) -> str:
    m, out = s["meeting"], []
    out.append(f"# 会议纪要:{m['topic']}\n")
    if m.get("description"):
        out.append(f"> {m['description']}\n")
    out.append(f"**参会专家**:{ '、'.join(e['name'] for e in s['experts']) }\n")
    if s.get("opening"):
        out.append(f"**主持人开场**:{s['opening']}\n")
    out.append("## 一、议程回顾\n")
    for i, a in enumerate(s["agenda"]):
        out.append(f"{i+1}. {a}")
    out.append("")
    out.append("## 二、各议题讨论摘要\n")
    for i, tp in enumerate(s["topics"]):
        out.append(f"### 议题 {i+1}:{tp['title']}\n")
        for tn in tp["turns"]:
            tag = "(追问)" if tn["followup"] else ("(参会者发言)" if tn["human"] else "")
            out.append(f"- **{tn['name']}**{tag}:{tn['text']}")
        if tp["conflict"]:
            out.append(f"\n> ⚠ 分歧:{tp['conflict'].get('summary','')};追问:{tp['conflict'].get('question','')}")
        if tp["summary"]:
            ss = tp["summary"]
            if ss.get("consensus"):
                out.append(f"\n**共识**:{ '；'.join(ss['consensus']) }")
            if ss.get("divergence"):
                out.append(f"**分歧**:{ '；'.join(ss['divergence']) }")
        out.append("")
    cc = s["conclusion"] or {}
    out.append("## 三、最终结论\n")
    out.append(cc.get("conclusion", "（无）") + "\n")
    if cc.get("risks"):
        out.append("## 四、风险清单\n")
        out.append("| 风险 | 涉及议题 | 相关专家观点 |\n|---|---|---|")
        for r in cc["risks"]:
            out.append(f"| {r.get('risk','')} | {r.get('topic','')} | {r.get('expert_view','')} |")
        out.append("")
    if cc.get("actions"):
        out.append("## 五、行动建议\n")
        for a in cc["actions"]:
            out.append(f"- {a}")
    return "\n".join(out)


def to_docx(s: dict) -> bytes:
    from docx import Document
    doc = Document()
    m = s["meeting"]
    doc.add_heading(f"会议纪要:{m['topic']}", level=0)
    if m.get("description"):
        doc.add_paragraph(m["description"])
    doc.add_paragraph("参会专家:" + "、".join(e["name"] for e in s["experts"]))
    doc.add_heading("一、议程回顾", level=1)
    for i, a in enumerate(s["agenda"]):
        doc.add_paragraph(f"{i+1}. {a}")
    doc.add_heading("二、各议题讨论摘要", level=1)
    for i, tp in enumerate(s["topics"]):
        doc.add_heading(f"议题 {i+1}:{tp['title']}", level=2)
        for tn in tp["turns"]:
            tag = "(追问)" if tn["followup"] else ("(参会者)" if tn["human"] else "")
            doc.add_paragraph(f"{tn['name']}{tag}:{tn['text']}")
        if tp["conflict"]:
            doc.add_paragraph(f"分歧:{tp['conflict'].get('summary','')};追问:{tp['conflict'].get('question','')}")
        if tp["summary"]:
            ss = tp["summary"]
            if ss.get("consensus"):
                doc.add_paragraph("共识:" + "；".join(ss["consensus"]))
            if ss.get("divergence"):
                doc.add_paragraph("分歧:" + "；".join(ss["divergence"]))
    cc = s["conclusion"] or {}
    doc.add_heading("三、最终结论", level=1)
    doc.add_paragraph(cc.get("conclusion", "（无）"))
    if cc.get("risks"):
        doc.add_heading("四、风险清单", level=1)
        tbl = doc.add_table(rows=1, cols=3); tbl.style = "Light Grid Accent 1"
        h = tbl.rows[0].cells; h[0].text, h[1].text, h[2].text = "风险", "涉及议题", "相关专家观点"
        for r in cc["risks"]:
            c = tbl.add_row().cells
            c[0].text, c[1].text, c[2].text = r.get("risk", ""), r.get("topic", ""), r.get("expert_view", "")
    if cc.get("actions"):
        doc.add_heading("五、行动建议", level=1)
        for a in cc["actions"]:
            doc.add_paragraph(a, style="List Bullet")
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


def _register_cjk() -> str | None:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    for path in _CJK_FONTS:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont("CJK", path, subfontIndex=0))
                return "CJK"
            except Exception:  # noqa: BLE001
                continue
    return None


def to_pdf(s: dict) -> bytes:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    from reportlab.lib import colors

    font = _register_cjk() or "Helvetica"
    ss = getSampleStyleSheet()
    H0 = ParagraphStyle("H0", parent=ss["Title"], fontName=font, fontSize=18, leading=24)
    H1 = ParagraphStyle("H1", parent=ss["Heading1"], fontName=font, fontSize=14, leading=20)
    H2 = ParagraphStyle("H2", parent=ss["Heading2"], fontName=font, fontSize=12, leading=18)
    P = ParagraphStyle("P", parent=ss["Normal"], fontName=font, fontSize=10.5, leading=16, alignment=TA_LEFT)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm,
                            leftMargin=18 * mm, rightMargin=18 * mm)
    el = []
    m = s["meeting"]
    esc = lambda x: str(x).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    el.append(Paragraph(f"会议纪要:{esc(m['topic'])}", H0))
    el.append(Paragraph("参会专家:" + esc("、".join(e["name"] for e in s["experts"])), P))
    el.append(Spacer(1, 6))
    el.append(Paragraph("一、议程回顾", H1))
    for i, a in enumerate(s["agenda"]):
        el.append(Paragraph(f"{i+1}. {esc(a)}", P))
    el.append(Paragraph("二、各议题讨论摘要", H1))
    for i, tp in enumerate(s["topics"]):
        el.append(Paragraph(f"议题 {i+1}:{esc(tp['title'])}", H2))
        for tn in tp["turns"]:
            tag = "(追问)" if tn["followup"] else ("(参会者)" if tn["human"] else "")
            el.append(Paragraph(f"<b>{esc(tn['name'])}</b>{tag}:{esc(tn['text'])}", P))
        if tp["summary"]:
            sm = tp["summary"]
            if sm.get("consensus"):
                el.append(Paragraph("共识:" + esc("；".join(sm["consensus"])), P))
            if sm.get("divergence"):
                el.append(Paragraph("分歧:" + esc("；".join(sm["divergence"])), P))
        el.append(Spacer(1, 4))
    cc = s["conclusion"] or {}
    el.append(Paragraph("三、最终结论", H1))
    el.append(Paragraph(esc(cc.get("conclusion", "（无）")), P))
    if cc.get("risks"):
        el.append(Paragraph("四、风险清单", H1))
        data = [["风险", "涉及议题", "相关专家观点"]]
        for r in cc["risks"]:
            data.append([Paragraph(esc(r.get("risk", "")), P), Paragraph(esc(r.get("topic", "")), P),
                         Paragraph(esc(r.get("expert_view", "")), P)])
        t = Table(data, colWidths=[60 * mm, 40 * mm, 65 * mm])
        t.setStyle(TableStyle([("FONTNAME", (0, 0), (-1, -1), font), ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                               ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cfc9ba")),
                               ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f3e7df")),
                               ("VALIGN", (0, 0), (-1, -1), "TOP")]))
        el.append(t)
    if cc.get("actions"):
        el.append(Paragraph("五、行动建议", H1))
        for a in cc["actions"]:
            el.append(Paragraph("• " + esc(a), P))
    doc.build(el)
    return buf.getvalue()
